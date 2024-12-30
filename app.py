import streamlit as st
st.set_page_config(page_title="Clanert Enviro Cloud", layout="wide")

import pandas as pd
from windrose import WindroseAxes
import matplotlib.pyplot as plt
import numpy as np
import io
import tempfile
import sqlite3
from streamlit_folium import st_folium
import folium
import calendar
from io import BytesIO
from matplotlib.dates import DateFormatter
import matplotlib.ticker as ticker
from datetime import time
from IPython.display import display, HTML
import os
from folium.plugins import HeatMap
import streamlit as st
from streamlit_folium import folium_static
from authlib.integrations.requests_client import OAuth2Session




# Google OAuth configuration
CLIENT_ID = "83123074595-m31sclpphtpej7k10pq6n3hfrh3oa0ik.apps.googleusercontent.com"
CLIENT_SECRET = "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"
REDIRECT_URI = "http://localhost:8501"

AUTHORIZATION_URL = "https://accounts.google.com/o/oauth2/auth"
TOKEN_URL = "https://oauth2.googleapis.com/token"

SCOPES = ["https://www.googleapis.com/auth/userinfo.email", "https://www.googleapis.com/auth/userinfo.profile"]

# Database setup
DB_NAME = "users.db"


def init_db():
    """Initialize the database."""
    conn = sqlite3.connect(DB_NAME)
    cursor = conn.cursor()
    
    # Create the table if it doesn't exist
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            email TEXT UNIQUE,
            name TEXT,
            password TEXT, -- For manual registration
            picture TEXT   -- For Google login
        )
    """)
    
    # Add 'password' column if it doesn't exist
    cursor.execute("PRAGMA table_info(users)")
    columns = [col[1] for col in cursor.fetchall()]
    if "password" not in columns:
        cursor.execute("ALTER TABLE users ADD COLUMN password TEXT")
    
    conn.commit()
    conn.close()



def save_user_to_db(email, name, password=None, picture=None):
    """Save the user to the database."""
    conn = sqlite3.connect(DB_NAME)
    cursor = conn.cursor()
    cursor.execute(
        "INSERT OR IGNORE INTO users (email, name, password, picture) VALUES (?, ?, ?, ?)",
        (email, name, password, picture),
    )
    conn.commit()
    conn.close()


def is_user_registered(email, password=None):
    """Check if the user is registered."""
    conn = sqlite3.connect(DB_NAME)
    cursor = conn.cursor()
    if password:  # Manual login
        cursor.execute("SELECT * FROM users WHERE email = ? AND password = ?", (email, password))
    else:  # Google login
        cursor.execute("SELECT * FROM users WHERE email = ?", (email,))
    user = cursor.fetchone()
    conn.close()
    return user


def google_login():
    """Initiate Google login."""
    oauth = OAuth2Session(client_id=CLIENT_ID, redirect_uri=REDIRECT_URI, scope=SCOPES)
    authorization_url, state = oauth.create_authorization_url(AUTHORIZATION_URL)
    st.session_state["state"] = state
    st.markdown(f"[Click here to log in with Google]({authorization_url})")


def fetch_user_info(token):
    """Fetch user information from Google."""
    oauth = OAuth2Session(client_id=CLIENT_ID, token=token)
    user_info = oauth.get("https://www.googleapis.com/oauth2/v3/userinfo").json()
    return user_info


# Initialize the database
init_db()

# Authentication
if "user" not in st.session_state:
    st.title("Welcome to Clanert Enviro Cloud (CEC model)")
    st.subheader("Please log in or register to continue")

    # Create tabs for Login and Registration
    tab1, tab2 = st.tabs(["Login", "Register"])

    # Login Tab
    with tab1:
        st.subheader("Login")
        login_method = st.radio("Choose your login method:", ["Manual Login", "Google Login"])

        if login_method == "Manual Login":
            email = st.text_input("Email")
            password = st.text_input("Password", type="password")
            if st.button("Login"):
                user = is_user_registered(email, password)
                if user:
                    st.session_state["user"] = {"email": user[1], "name": user[2], "picture": user[4]}
                    st.success(f"Welcome back, {user[2]}!")
                    st.rerun()
                else:
                    st.error("Invalid email or password. Please try again.")
        elif login_method == "Google Login":
            google_login()
            query_params = st.query_params
            code = query_params.get("code", None)
            state = query_params.get("state", None)

            if code:
                if state == st.session_state.get("state"):
                    oauth = OAuth2Session(client_id=CLIENT_ID, client_secret=CLIENT_SECRET, redirect_uri=REDIRECT_URI)
                    token = oauth.fetch_token(
                        TOKEN_URL,
                        authorization_response=f"{REDIRECT_URI}?code={code}",
                        client_secret=CLIENT_SECRET,
                    )
                    st.session_state["token"] = token
                    user_info = fetch_user_info(token)
                    st.session_state["user"] = user_info

                    # Check if the user is registered
                    if not is_user_registered(user_info["email"]):
                        save_user_to_db(user_info["email"], user_info["name"], picture=user_info["picture"])
                        st.success("Registration successful! Welcome to ClanertWAA.")
                    else:
                        st.success("Login successful! Welcome back.")
                    st.rerun()
                else:
                    st.error("State mismatch. Please try again.")

    # Registration Tab
    with tab2:
        st.subheader("Register")
        name = st.text_input("Full Name")
        email = st.text_input("Email (Registration)")
        password = st.text_input("Password (Registration)", type="password")
        if st.button("Register"):
            if name and email and password:
                if not is_user_registered(email):
                    save_user_to_db(email, name, password=password)
                    st.success("Registration successful! Please log in.")
                else:
                    st.error("This email is already registered. Please log in.")
            else:
                st.error("Please fill in all fields.")

    # Stop the app until user logs in
    st.stop()


# Display user info in the sidebar
user_info = st.session_state["user"]
st.sidebar.success(f"Welcome, {user_info.get('name', user_info['email'])}!")

# Validate the picture URL
picture_url = user_info.get("picture", None)

if picture_url:
    try:
        # Attempt to display the user's picture
        st.sidebar.image(picture_url, width=100)
    except Exception as e:
        # If there's an error, display a placeholder
        st.sidebar.image("https://i.ibb.co/Z8VjzNV/CLANERT-SUSTAIN-logo-fin-AL.png", width=100, caption="Profile Picture")
else:
    # Display a placeholder if no picture is available
    st.sidebar.image("https://i.ibb.co/Z8VjzNV/CLANERT-SUSTAIN-logo-fin-AL.png", width=100, caption="No Profile Picture")

# Footer
st.markdown("---")
st.markdown(
    """
    <p style="text-align: center; font-size: 14px; font-variant: small-caps;">
        <strong> all rights reserved and copyright. this is owned by clanert sustain company ltd.</strong> further details can be found at 
        <a href="https://www.clanertsustain.co.tz" target="_blank" style="color: #FF5722; text-decoration: none;">
            🌐 www.clanertsustain.co.tz
        </a> or contact us via email at 
        <a href="mailto:info@clanertsustain.co.tz" style="color: #FF5722; text-decoration: none;">
            ✉️ info@clanertsustain.co.tz
        </a>.
    </p>
    """,
    unsafe_allow_html=True,
)



# Add your main app code here
st.title("🌬️ Clanert Enviro Cloud")
st.write("Main application content starts here...")


# Sidebar navigation
st.sidebar.title("Navigation")
options = st.sidebar.radio(
    "Choose a section:",
    ["Instructions", "Wind Analysis", "Air Quality Analysis", "Multivariable Analysis", "Plume Dispersion Model", "About"]
)

# Instructions section
if options == "Instructions":
    st.subheader("📖 Instructions")
    st.write("""
### General Instructions:
- Log in to access the dashboard. You can log in using your email and password or with Google authentication.
- Use the sidebar navigation to explore different sections of the dashboard.

### For Wind Analysis:
- **Upload a CSV file** containing:
  - **MO**: Month (1-12)
  - **WD10M**: Wind Direction (0-360°)
  - **WS10M**: Wind Speed (m/s)
- **Features**:
  - Select specific months to filter the data.
  - Visualize wind data using:
    - **Wind Class Distribution**: A bar chart showing wind speed categories.
    - **Wind Rose Visualization**: A circular chart showing wind direction and speed.
  - Overlay wind rose charts on geographical maps using latitude and longitude inputs.
- **Download Options**:
  - Save Wind Class and Wind Rose charts as PNG files.

### For Air Quality Analysis:
- **Upload one or two CSV files** containing air quality data, including:
  - **Pollutants**: Columns for **O3**, **NO2**, **CO**, **SO2**, **PM10**, and **PM2.5**.
  - **Duration**: A column in HH:MM:SS format representing time intervals.
- **Features**:
  - View pollutant statistics (mean, max, min, and standard deviation).
  - Calculate Air Quality Index (AQI) for PM10 and PM2.5, and display AQI categories.
  - Compare pollutant concentrations between two locations.
  - Visualize pollutant trends over time using line plots.
- **Download Options**:
  - Save pollutant concentration plots and AQI charts as PNG files.

### For Multivariable Analysis:
- **Upload CSV files** for multiple locations to compare pollutant trends.
- Analyze specific time periods by setting start and end times.
- Visualize and compare:
  - PM10 and PM2.5 concentrations across multiple locations.
- View a summary table with AQI values, categories, and color-coded visual indicators.

### For Report Generation:
- Generate a detailed air quality report in Word format.
- The report includes:
  - Comparisons with TBS/WHO standards.
  - AQI-based recommendations for each location.
  - Summarized pollutant data and AQI analysis.
- Download the report for sharing or record-keeping.

### Notes:
- Ensure all CSV files match the required format.
- If errors occur, double-check the column names and data structure.
- For help, refer to the "About" section or contact support.
    """)


# Wind Analysis section
if options == "Wind Analysis":
    st.subheader(" Wind Analysis Dashboard")
    
    # Upload file for wind data
    uploaded_file = st.sidebar.file_uploader("📂 Upload your wind data CSV", type=["csv"])

    if uploaded_file:
        data = pd.read_csv(uploaded_file)

        # Validate columns
        if not all(col in data.columns for col in ["MO", "WD10M", "WS10M"]):
            st.error("The uploaded file must contain 'MO', 'WD10M', and 'WS10M' columns.")
        else:
            # Month selection
            selected_months = st.sidebar.multiselect(
                "📅 Select Months",
                options=list(range(1, 13)),
                format_func=lambda x: f"{x} - {calendar.month_name[x]}"
            )

            if not selected_months:
                st.warning("Please select at least one month.")
            else:
                filtered_data = data[data['MO'].isin(selected_months)]

                if filtered_data.empty:
                    st.warning(f"No data available for the selected months: {selected_months}.")
                else:
                    # Layout for visualizations
                    col1, col2 = st.columns(2)

                    # Wind Class Distribution
                    with col1:
                        st.subheader("📊 Wind Class Distribution")
                        bins = [0, 2, 4, 6, 8, 10, 12, 15, 20, float('inf')]
                        labels = ['0-2', '2-4', '4-6', '6-8', '8-10', '10-12', '12-15', '15-20', '20+']

                        filtered_data['Wind_Class'] = pd.cut(
                            filtered_data['WS10M'], bins=bins, labels=labels, right=False
                        )

                        wind_class_distribution = (
                            filtered_data['Wind_Class']
                            .value_counts(normalize=True)
                            .sort_index() * 100
                        )

                        wind_class_df = wind_class_distribution.reset_index()
                        wind_class_df.columns = ['Wind Class', 'Percentage']

                        fig_class, ax_class = plt.subplots(figsize=(5, 5))
                        ax_class.bar(
                            wind_class_df['Wind Class'],
                            wind_class_df['Percentage'],
                            color='skyblue', edgecolor='black'
                        )
                        ax_class.set_xlabel('Wind Class (m/s)')
                        ax_class.set_ylabel('Percentage (%)')
                        ax_class.set_title('Wind Class Distribution')

                        st.pyplot(fig_class)

                        # Download button
                        buffer_class = io.BytesIO()
                        fig_class.savefig(buffer_class, format="png")
                        buffer_class.seek(0)
                        st.download_button(
                            "📥 Download Wind Class Chart",
                            data=buffer_class,
                            file_name="wind_class_distribution.png",
                            mime="image/png"
                        )

                    # Wind Rose Visualization
                with col2:
                    st.subheader(" Wind Rose Visualization")

                    def plot_wind_rose_with_separated_legend(data_subset):
                        # Calculate resultant wind vector
                        u = -data_subset['WS10M'] * np.sin(np.radians(data_subset['WD10M']))
                        v = -data_subset['WS10M'] * np.cos(np.radians(data_subset['WD10M']))
                        resultant_speed = np.sqrt(u.mean()**2 + v.mean()**2)
                        resultant_direction = np.degrees(np.arctan2(u.mean(), v.mean()))
                        if resultant_direction < 0:
                            resultant_direction += 360

                        # Create the wind rose plot
                        fig, ax = plt.subplots(figsize=(6, 6), subplot_kw={'projection': 'windrose'})
                        ax.bar(
                            data_subset['WD10M'], data_subset['WS10M'],
                            normed=True, bins=[0, 2, 4, 6, 8, 10, 12, 15, 20],
                            edgecolor='black'
                        )

                        # Add resultant vector arrow
                        ax.arrow(
                            resultant_direction, 0, 0, ax.get_rmax() * 0.9,
                            color="red", lw=1, alpha=0.8, length_includes_head=True,
                            head_width=ax.get_rmax() * 0.01, head_length=ax.get_rmax() * 0.02
                        )

                        # Add resultant wind information
                        ax.text(
                            0.5, -0.1,
                            f"Resultant: {resultant_speed:.2f} m/s, Dir: {resultant_direction:.0f}°",
                            transform=ax.transAxes, fontsize=10, color="red", ha="center"
                        )

                        # Position the legend to the right of the wind rose
                        legend = ax.legend(
                            title="Wind Speed (m/s)",
                            loc="upper left",
                            bbox_to_anchor=(1.05, 1),
                            fontsize=8,
                            title_fontsize=10
                        )

                        # Adjust the layout to avoid clipping
                        fig.tight_layout()

                        return fig

                    fig_wind_rose = plot_wind_rose_with_separated_legend(filtered_data)
                    st.pyplot(fig_wind_rose)

                    # Download Button
                    buffer_rose = io.BytesIO()
                    fig_wind_rose.savefig(buffer_rose, format="png")
                    buffer_rose.seek(0)
                    st.download_button(
                        "📥 Download Wind Rose Chart",
                        data=buffer_rose,
                        file_name="wind_rose.png",
                        mime="image/png"
                    )

                    # Map Visualization
                    st.subheader("🗺️ Wind Rose Overlay on Map")
                    lat = st.number_input("Latitude:", value=0.0, format="%.6f")
                    lon = st.number_input("Longitude:", value=0.0, format="%.6f")

                    map_layer = st.radio(
                        "Choose Map Layer:",
                        options=["OpenStreetMap", "Esri Satellite Hybrid"]
                    )

                    def save_clean_wind_rose(data_subset):
                        ax = WindroseAxes.from_ax()
                        ax.bar(
                            data_subset['WD10M'], data_subset['WS10M'],
                            normed=True, bins=[0, 2, 4, 6, 8, 10, 12, 15, 20],
                            edgecolor="black"
                        )
                        ax.grid(False)
                        ax.set_yticks([])
                        ax.set_xticks([])
                        ax.spines["polar"].set_visible(False)
                        buffer = io.BytesIO()
                        plt.savefig(buffer, format="png", dpi=150, bbox_inches="tight", transparent=True)
                        plt.close()
                        buffer.seek(0)
                        return buffer

                    windrose_buffer = save_clean_wind_rose(filtered_data)

                    with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp_file:
                        tmp_file.write(windrose_buffer.getvalue())
                        temp_image_path = tmp_file.name

                    if map_layer == "OpenStreetMap":
                        map_plot = folium.Map(location=[lat, lon], zoom_start=14, tiles="OpenStreetMap")
                    else:
                        map_plot = folium.Map(
                            location=[lat, lon],
                            zoom_start=20,
                            tiles="https://services.arcgisonline.com/arcgis/rest/services/World_Imagery/MapServer/tile/{z}/{y}/{x}",
                            attr="Esri, USGS | Source: Esri, DigitalGlobe, GeoEye, Earthstar Geographics, CNES/Airbus DS, USDA, USGS, AeroGRID, IGN, and the GIS User Community"
                        )

                    folium.raster_layers.ImageOverlay(
                        image=temp_image_path,
                        bounds=[[lat - 0.001, lon - 0.001], [lat + 0.001, lon + 0.001]],
                        opacity=0.8
                    ).add_to(map_plot)

                    st_folium(map_plot, width=900, height=500)

# airquality analysis
if options == "Air Quality Analysis":
    st.subheader("📉 Air Quality Analysis Dashboard")
   
  
    # Upload files for air quality data
    uploaded_file1 = st.sidebar.file_uploader("Upload the first air quality CSV file:", type=["csv"], key="aqi_file_1")
    uploaded_file2 = st.sidebar.file_uploader("Upload the second air quality CSV file (optional):", type=["csv"], key="aqi_file_2")

    # Code continues...
    if uploaded_file1:
        # Input for the first location
        sampling_location1 = st.text_input("Enter the name of the first sampling location:", "Location 1")
        data1 = pd.read_csv(uploaded_file1)
        data1['Duration'] = pd.to_datetime(data1['Duration'], format='%H:%M:%S')

        # AQI calculations for the first location
        def calculate_aqi(concentration, breakpoints):
            for bp in breakpoints:
                if bp[0] <= concentration <= bp[1]:
                    aqi = (bp[3] - bp[2]) / (bp[1] - bp[0]) * (concentration - bp[0]) + bp[2]
                    return round(aqi)
            return None

        def get_aqi_category(aqi):
            if 0 <= aqi <= 50:
                return 'Good', '#00e400'
            elif 51 <= aqi <= 100:
                return 'Moderate', '#ffff00'
            elif 101 <= aqi <= 150:
                return 'Unhealthy for Sensitive Groups', '#ff7e00'
            elif 151 <= aqi <= 200:
                return 'Unhealthy', '#ff0000'
            elif 201 <= aqi <= 300:
                return 'Very Unhealthy', '#99004c'
            elif 301 <= aqi <= 500:
                return 'Hazardous', '#7e0023'
            else:
                return 'Out of Range', '#ffffff'

        # Define AQI breakpoints for PM10 and PM2.5
        pm10_breakpoints = [
            [0, 54, 0, 50],
            [55, 154, 51, 100],
            [155, 254, 101, 150],
            [255, 354, 151, 200],
            [355, 424, 201, 300],
            [425, 504, 301, 400],
            [505, 604, 401, 500]
        ]

        pm25_breakpoints = [
            [0, 12, 0, 50],
            [12.1, 35.4, 51, 100],
            [35.5, 55.4, 101, 150],
            [55.5, 150.4, 151, 200],
            [150.5, 250.4, 201, 300],
            [250.5, 350.4, 301, 400],
            [350.5, 500.4, 401, 500]
        ]

        # WHO standards for pollutants (µg/m³)
        tbs_standards = {
            'PM10': 50,      # WHO daily guideline for PM10
            'PM2.5': 25,     # WHO daily guideline for PM2.5
            'O3': 100,       # WHO 8-hour guideline
            'NO2': 40,       # WHO annual guideline
            'CO': 4 * 1000,  # WHO 8-hour guideline in mg/m³ converted to µg/m³
            'SO2': 20        # WHO 24-hour guideline
        }

        # Add AQI columns
        data1['PM10_AQI'] = data1['PM10'].apply(lambda x: calculate_aqi(x, pm10_breakpoints))
        data1['PM2.5_AQI'] = data1['PM2.5'].apply(lambda x: calculate_aqi(x, pm25_breakpoints))

        # Display statistics for the first location
        st.subheader(f"📊 Statistics for {sampling_location1}")
        pollutants = ['O3', 'NO2', 'CO', 'SO2', 'PM10', 'PM2.5']
        stats1 = {pollutant: {'mean': data1[pollutant].mean(), 'max': data1[pollutant].max(),
                              'min': data1[pollutant].min(), 'std': data1[pollutant].std()}
                  for pollutant in pollutants}

        for pollutant, values in stats1.items():
            st.write(f"{pollutant}: Mean = {values['mean']:.2f}, Max = {values['max']:.2f}, "
                     f"Min = {values['min']:.2f}, Std Dev = {values['std']:.2f}")

        # AQI summary table
        st.subheader("📋 AQI Summary Table")
        aqi_table1 = pd.DataFrame({
            'Pollutant': ['PM10', 'PM2.5'],
            'Mean AQI': [data1['PM10_AQI'].mean(), data1['PM2.5_AQI'].mean()],
            'Max AQI': [data1['PM10_AQI'].max(), data1['PM2.5_AQI'].max()],
            'Min AQI': [data1['PM10_AQI'].min(), data1['PM2.5_AQI'].min()]
        })
        aqi_table1['Category'], aqi_table1['Color'] = zip(*aqi_table1['Mean AQI'].apply(get_aqi_category))
        st.dataframe(aqi_table1)

        # Line plot for PM10 and PM2.5
        st.subheader(f"📈 PM10 and PM2.5 Concentrations at {sampling_location1}")
        fig1, ax1 = plt.subplots(figsize=(12, 8))
        ax1.plot(data1['Duration'], data1['PM10'], label='PM10', color='blue', linewidth=2)
        ax1.plot(data1['Duration'], data1['PM2.5'], label='PM2.5', color='green', linewidth=2)
        ax1.axhline(y=stats1['PM10']['mean'], color='blue', linestyle='--', label=f'PM10 Mean ({stats1["PM10"]["mean"]:.2f})')
        ax1.axhline(y=stats1['PM2.5']['mean'], color='green', linestyle='--', label=f'PM2.5 Mean ({stats1["PM2.5"]["mean"]:.2f})')
        ax1.set_xlabel('Duration (HH:MM:SS)')
        ax1.set_ylabel('Concentration (µg/m³)')
        ax1.set_title(f"PM10 and PM2.5 Concentrations at {sampling_location1}")
        ax1.legend()
        # Apply DateFormatter to format the x-axis as HH:MM:SS
        time_formatter = DateFormatter('%H:%M:%S')
        ax1.xaxis.set_major_formatter(time_formatter)
        plt.xticks(rotation=45)
        st.pyplot(fig1)
       

         # Line plot for PM10 and PM2.5 with WHO standards
        st.subheader(f"📈 PM10 and PM2.5 Concentrations at {sampling_location1}")
        fig1, ax1 = plt.subplots(figsize=(12, 6))
        ax1.plot(data1['Duration'], data1['PM10'], label='PM10', color='blue', linewidth=2)
        ax1.plot(data1['Duration'], data1['PM2.5'], label='PM2.5', color='green', linewidth=2)
        ax1.axhline(y=tbs_standards['PM10'], color='purple', linestyle='--', label='PM10 TBS Guideline (50 µg/m³)')
        ax1.axhline(y=tbs_standards['PM2.5'], color='black', linestyle='--', label='PM2.5 TBS Guideline (25 µg/m³)')
        ax1.set_xlabel('Duration (HH:MM:SS)')
        ax1.set_ylabel('Concentration (µg/m³)')
        ax1.set_title(f"PM10 and PM2.5 Concentrations at {sampling_location1}")
        ax1.legend()
        

        # Apply DateFormatter to format the x-axis as HH:MM:SS
        time_formatter = DateFormatter('%H:%M:%S')
        ax1.xaxis.set_major_formatter(time_formatter)
        plt.xticks(rotation=45)
        st.pyplot(fig1)

       

        # Download button for the plot
        buf1 = BytesIO()
        fig1.savefig(buf1, format="png")
        buf1.seek(0)
        st.download_button(
            label="📥 Download PM10 and PM2.5 Plot",
            data=buf1,
            file_name="pm10_pm25_plot.png",
            mime="image/png"
        )

      
        # Plot other pollutants
        for pollutant in ['O3', 'NO2', 'CO', 'SO2']:
            st.subheader(f"📉 {pollutant} Concentration at {sampling_location1}")
            fig2, ax2 = plt.subplots(figsize=(12, 6))
            ax2.plot(data1['Duration'], data1[pollutant], label=pollutant, color='green', linewidth=2)
            ax2.axhline(y=tbs_standards[pollutant], color='orange', linestyle='--',
                        label=f'{pollutant} TBS Guideline ({tbs_standards[pollutant]} µg/m³)')
            ax2.set_xlabel('Duration (HH:MM:SS)')
            ax2.set_ylabel('Concentration (µg/m³)')
            ax2.set_title(f"{pollutant} Concentration at {sampling_location1}")
            ax2.legend()
            
            # Apply DateFormatter to format the x-axis as HH:MM:SS
            time_formatter = DateFormatter('%H:%M:%S')
            ax2.xaxis.set_major_formatter(time_formatter)

            # Rotate x-axis labels for better readability
            plt.xticks(rotation=45)
            st.pyplot(fig2)

            # Download button for the plot
        buf2 = BytesIO()
        fig2.savefig(buf2, format="png")
        buf2.seek(0)
        st.download_button(
            label="📥 Download SO2",
            data=buf2,
            file_name="Gas_emissions.png",
            mime="image/png"
        )


    # Second location analysis (if file uploaded)
    if uploaded_file2:
        sampling_location2 = st.text_input("Enter the name of the second sampling location:", "Location 2")
        data2 = pd.read_csv(uploaded_file2)
        data2['Duration'] = pd.to_datetime(data2['Duration'], format='%H:%M:%S')

        # AQI calculations for the second location
        data2['PM10_AQI'] = data2['PM10'].apply(lambda x: calculate_aqi(x, pm10_breakpoints))
        data2['PM2.5_AQI'] = data2['PM2.5'].apply(lambda x: calculate_aqi(x, pm25_breakpoints))

        # Comparison plot for PM10 and PM2.5
        st.subheader(f"🔍 PM10 and PM2.5 Comparison: {sampling_location1} vs {sampling_location2}")
        fig3, ax3 = plt.subplots(figsize=(12, 8))
        ax3.plot(data1['Duration'], data1['PM10'], label=f'PM10 at {sampling_location1}', color='blue')
        ax3.plot(data2['Duration'], data2['PM10'], label=f'PM10 at {sampling_location2}', color='cyan')
        ax3.plot(data1['Duration'], data1['PM2.5'], label=f'PM2.5 at {sampling_location1}', color='red')
        ax3.plot(data2['Duration'], data2['PM2.5'], label=f'PM2.5 at {sampling_location2}', color='orange')
        ax3.set_title(f"PM10 and PM2.5 Concentration Comparison")
        ax3.legend()
        st.pyplot(fig3)
# Apply DateFormatter to format the x-axis as HH:MM:SS
        time_formatter = DateFormatter('%H:%M:%S')
        ax1.xaxis.set_major_formatter(time_formatter)

        # Rotate x-axis labels for better readability
        plt.xticks(rotation=45)
        st.pyplot(fig1)


# Function to calculate AQI
def calculate_aqi(concentration, breakpoints):
    for bp in breakpoints:
        if bp[0] <= concentration <= bp[1]:
            aqi = (bp[3] - bp[2]) / (bp[1] - bp[0]) * (concentration - bp[0]) + bp[2]
            return round(aqi)
    return None

# Function to get AQI category and color
def get_aqi_category(aqi):
    if 0 <= aqi <= 50:
        return 'Good', '#00e400'
    elif 51 <= aqi <= 100:
        return 'Moderate', '#ffff00'
    elif 101 <= aqi <= 150:
        return 'Unhealthy for Sensitive Groups', '#ff7e00'
    elif 151 <= aqi <= 200:
        return 'Unhealthy', '#ff0000'
    elif 201 <= aqi <= 300:
        return 'Very Unhealthy', '#99004c'
    elif 301 <= aqi <= 500:
        return 'Hazardous', '#7e0023'
    else:
        return 'Out of Range', '#ffffff'

# AQI breakpoints for PM10 and PM2.5
pm10_breakpoints = [
    [0, 54, 0, 50],
    [55, 154, 51, 100],
    [155, 254, 101, 150],
    [255, 354, 151, 200],
    [355, 424, 201, 300],
    [425, 504, 301, 400],
    [505, 604, 401, 500]
]

pm25_breakpoints = [
    [0, 12, 0, 50],
    [12.1, 35.4, 51, 100],
    [35.5, 55.4, 101, 150],
    [55.5, 150.4, 151, 200],
    [150.5, 250.4, 201, 300],
    [250.5, 350.4, 301, 400],
    [350.5, 500.4, 401, 500]
]

# MULTIVARIABLE ANALYSIS
if options == "Multivariable Analysis":
    st.title("Multivariable Air Quality Analysis")

        # Function to calculate AQI
    def calculate_aqi(concentration, breakpoints):
        for bp in breakpoints:
            if bp[0] <= concentration <= bp[1]:
                aqi = (bp[3] - bp[2]) / (bp[1] - bp[0]) * (concentration - bp[0]) + bp[2]
                return round(aqi)
        return None

    # AQI breakpoints for PM10 and PM2.5
    pm10_breakpoints = [
        [0, 54, 0, 50],
        [55, 154, 51, 100],
        [155, 254, 101, 150],
        [255, 354, 151, 200],
        [355, 424, 201, 300],
        [425, 504, 301, 400],
        [505, 604, 401, 500]
    ]

    pm25_breakpoints = [
        [0, 12, 0, 50],
        [12.1, 35.4, 51, 100],
        [35.5, 55.4, 101, 150],
        [55.5, 150.4, 151, 200],
        [150.5, 250.4, 201, 300],
        [250.5, 350.4, 301, 400],
        [350.5, 500.4, 401, 500]
    ]



    # Number of locations to analyze
    num_locations = st.number_input("Enter the number of locations to analyze:", min_value=1, step=1)

    file_uploads = []
    sampling_locations = []

    # File upload and location names
    for i in range(num_locations):
        file = st.file_uploader(f"Upload CSV file for location {i+1}:", type="csv", key=f"file_{i}")
        location_name = st.text_input(f"Enter the name for sampling location {i+1}:", key=f"location_{i}")
        if file and location_name:
            file_uploads.append((file, location_name))

    # Data processing and plotting
    if file_uploads:
        dataframes = []
        for file, location in file_uploads:
            df = pd.read_csv(file)
            df['Duration'] = pd.to_datetime(df['Duration'], format='%H:%M:%S')
            df['PM10_AQI'] = df['PM10'].apply(lambda x: calculate_aqi(x, pm10_breakpoints))
            df['PM2.5_AQI'] = df['PM2.5'].apply(lambda x: calculate_aqi(x, pm25_breakpoints))
            dataframes.append((location, df))

        # Time selection
        time_analysis = st.checkbox("Analyze specific time periods?")
        if time_analysis:
            start_time = st.time_input("Start time:")
            end_time = st.time_input("End time:")

            def filter_time_period(df):
                return df[(df['Duration'].dt.time >= start_time) & (df['Duration'].dt.time <= end_time)]
            dataframes = [(location, filter_time_period(df)) for location, df in dataframes]

        # Plotting PM10
        st.subheader("PM10 Concentration Comparison Across Locations")
        fig, ax = plt.subplots(figsize=(14, 10))
        for location, df in dataframes:
            ax.plot(df['Duration'], df['PM10'], label=f'PM10 at {location}', linewidth=2)
            ax.axhline(y=df['PM10'].mean(), linestyle='--', label=f'{location} PM10 Mean ({df["PM10"].mean():.2f})')
        ax.set_xlabel('Duration (HH:MM:SS)', fontsize=12)
        ax.set_ylabel('Concentration (µg/m³)', fontsize=12)
        ax.set_title('PM10 Concentration Comparison Across Locations', fontsize=14, fontweight='bold')
        ax.legend(fontsize=10)
        ax.grid(visible=True, linestyle='--', alpha=0.6)
        ax.xaxis.set_major_formatter(DateFormatter('%H:%M:%S'))
        plt.xticks(rotation=45, fontsize=10)
        plt.yticks(fontsize=10)
        st.pyplot(fig)

        # Plotting PM2.5
        st.subheader("PM2.5 Concentration Comparison Across Locations")
        fig, ax = plt.subplots(figsize=(14, 10))
        for location, df in dataframes:
            ax.plot(df['Duration'], df['PM2.5'], label=f'PM2.5 at {location}', linewidth=2)
            ax.axhline(y=df['PM2.5'].mean(), linestyle='--', label=f'{location} PM2.5 Mean ({df["PM2.5"].mean():.2f})')
        ax.set_xlabel('Duration (HH:MM:SS)', fontsize=12)
        ax.set_ylabel('Concentration (µg/m³)', fontsize=12)
        ax.set_title('PM2.5 Concentration Comparison Across Locations', fontsize=14, fontweight='bold')
        ax.legend(fontsize=10)
        ax.grid(visible=True, linestyle='--', alpha=0.6)
        ax.xaxis.set_major_formatter(DateFormatter('%H:%M:%S'))
        plt.xticks(rotation=45, fontsize=10)
        plt.yticks(fontsize=10)
        st.pyplot(fig)

            # Generate a color-coded AQI comparison table
        st.subheader("Air Quality Index (AQI) Comparison Table")

        # Function to generate AQI category and color
        def get_aqi_category_and_color(aqi_value):
            if aqi_value <= 50:
                return "Good", "#00E400"
            elif aqi_value <= 100:
                return "Moderate", "#FFFF00"
            elif aqi_value <= 150:
                return "Unhealthy for Sensitive Groups", "#FF7E00"
            elif aqi_value <= 200:
                return "Unhealthy", "#FF0000"
            elif aqi_value <= 300:
                return "Very Unhealthy", "#8F3F97"
            else:
                return "Hazardous", "#7E0023"

        # Collect summary data for each location
        summary_data = []

        for location, df in dataframes:
            avg_pm10_aqi = df['PM10_AQI'].mean()
            avg_pm25_aqi = df['PM2.5_AQI'].mean()

            pm10_category, pm10_color = get_aqi_category_and_color(avg_pm10_aqi)
            pm25_category, pm25_color = get_aqi_category_and_color(avg_pm25_aqi)

            summary_data.append({
                'Location': location,
                'Avg PM10 AQI': round(avg_pm10_aqi, 2),
                'PM10 Category': pm10_category,
                'PM10 Color': pm10_color,
                'Avg PM2.5 AQI': round(avg_pm25_aqi, 2),
                'PM2.5 Category': pm25_category,
                'PM2.5 Color': pm25_color
            })

        # Convert summary data to DataFrame
        summary_df = pd.DataFrame(summary_data)

        # Render the table with color-coded AQI categories
        st.markdown("### Summary Table with Color-Coded AQI Categories")
        for _, row in summary_df.iterrows():
            st.markdown(
                f"""
                <div style='padding:10px; border:1px solid #ddd; margin-bottom:10px; border-radius:5px;'>
                    <h4 style='margin:0;'>{row['Location']}</h4>
                    <p><b>PM10 AQI:</b> {row['Avg PM10 AQI']} - 
                    <span style='color:black; background-color:{row['PM10 Color']}; 
                                    padding:3px 6px; border-radius:3px;'>{row['PM10 Category']}</span></p>
                    <p><b>PM2.5 AQI:</b> {row['Avg PM2.5 AQI']} - 
                    <span style='color:black; background-color:{row['PM2.5 Color']}; 
                                    padding:3px 6px; border-radius:3px;'>{row['PM2.5 Category']}</span></p>
                </div>
                """, unsafe_allow_html=True
            )

# Report Generator Section

        st.subheader("📄 Air Quality Report Generator")
        st.write("Generate a comprehensive report based on the processed data from the Multivariable Analysis section. This report includes comparisons with TBS/WHO standards, AQI-based recommendations, and color-coded tables.")

        # Check if `dataframes` exists
        if 'dataframes' in locals() and dataframes:
            # Function to compare concentrations with TBS/WHO standards
            def compare_with_standard(average, standard):
                return "Meets Standard" if average <= standard else "Exceeds Standard"

            # Generate AQI-based advice
            def generate_aqi_advice(aqi, category):
                if category == 'Good':
                    return "Air quality is excellent. No actions required."
                elif category == 'Moderate':
                    return "Air quality is acceptable. Sensitive groups should reduce prolonged outdoor activities."
                elif category == 'Unhealthy for Sensitive Groups':
                    return "Air quality is unhealthy for sensitive groups. Limit outdoor activities."
                elif category == 'Unhealthy':
                    return "Air quality is unhealthy. Everyone should limit outdoor activities."
                elif category == 'Very Unhealthy':
                    return "Air quality is very unhealthy. Stay indoors and use air purifiers."
                elif category == 'Hazardous':
                    return "Air quality is hazardous. Avoid all outdoor activities."
                else:
                    return "Air quality data is out of range. Please consult authorities."

            # Generate detailed Word report
            def generate_detailed_report(dataframes, file_name="air_quality_report.docx"):
                from docx import Document
                from datetime import datetime

                # Define the TBS/WHO limits inside the function
                TBS_PM10_LIMIT_24H = 50
                TBS_PM25_LIMIT_24H = 25


                doc = Document()
                doc.add_heading('Air Quality Analysis Report', level=1)
                doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d')}")
                doc.add_paragraph(
                    "This report provides a comprehensive analysis of air quality data, including comparisons with TBS/WHO standards, "
                    "AQI values, and actionable recommendations."
                )

                # Real concentration summary
                doc.add_heading('Real Concentration Analysis (TBS/WHO Standards)', level=2)
                table = doc.add_table(rows=1, cols=5)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = "Location"
                hdr_cells[1].text = "Average PM10 (µg/m³)"
                hdr_cells[2].text = "PM10 Meets TBS/WHO (50)?"
                hdr_cells[3].text = "Average PM2.5 (µg/m³)"
                hdr_cells[4].text = "PM2.5 Meets TBS/WHO (25)?"

                for location, df in dataframes:
                    avg_pm10 = df['PM10'].mean()
                    avg_pm25 = df['PM2.5'].mean()
                    cells = table.add_row().cells
                    cells[0].text = location
                    cells[1].text = f"{avg_pm10:.2f}"
                    cells[2].text = compare_with_standard(avg_pm10, TBS_PM10_LIMIT_24H)
                    cells[3].text = f"{avg_pm25:.2f}"
                    cells[4].text = compare_with_standard(avg_pm25, TBS_PM25_LIMIT_24H)

                # AQI analysis
                doc.add_heading('AQI Analysis and Recommendations', level=2)
                aqi_table = doc.add_table(rows=1, cols=6)
                aqi_table.style = 'Table Grid'
                hdr_cells = aqi_table.rows[0].cells
                hdr_cells[0].text = "Location"
                hdr_cells[1].text = "PM10 AQI"
                hdr_cells[2].text = "PM10 Category"
                hdr_cells[3].text = "PM2.5 AQI"
                hdr_cells[4].text = "PM2.5 Category"
                hdr_cells[5].text = "Advice"

                for location, df in dataframes:
                    avg_pm10_aqi = df['PM10_AQI'].mean()
                    avg_pm25_aqi = df['PM2.5_AQI'].mean()
                    pm10_category, _ = get_aqi_category(avg_pm10_aqi)
                    pm25_category, _ = get_aqi_category(avg_pm25_aqi)

                    advice = (
                        f"PM10: {generate_aqi_advice(avg_pm10_aqi, pm10_category)}\n"
                        f"PM2.5: {generate_aqi_advice(avg_pm25_aqi, pm25_category)}"
                    )

                    cells = aqi_table.add_row().cells
                    cells[0].text = location
                    cells[1].text = f"{avg_pm10_aqi:.2f}"
                    cells[2].text = pm10_category
                    cells[3].text = f"{avg_pm25_aqi:.2f}"
                    cells[4].text = pm25_category
                    cells[5].text = advice

                doc.save(file_name)
                return file_name

            # Generate Report Button
            if st.button("Generate Report"):
                report_file = generate_detailed_report(dataframes)
                st.success("Report generated successfully!")
                with open(report_file, "rb") as f:
                    st.download_button("Download Report", f, file_name="air_quality_report.docx")
        else:
            st.warning("No data available. Please complete the 'Multivariable Analysis' section first.")





if options == "Plume Dispersion Model":
    st.header("Plume Dispersion Model")

    # Step 1: Load Meteorology Data
    def load_meteorology_data(csv_path):
        """
        Load meteorological data from a CSV file.
        The file should contain columns: 'YEAR', 'MO', 'DY', 'HR', 'WS10M', 'WD10M', etc.

        Parameters:
        csv_path : str : Path to the meteorological CSV file

        Returns:
        DataFrame : Meteorological data
        """
        try:
            met_data = pd.read_csv(csv_path)
            required_columns = {'YEAR', 'MO', 'DY', 'HR', 'WS10M', 'WD10M'}
            if not required_columns.issubset(met_data.columns):
                st.error(f"CSV file must contain columns: {required_columns}")
                return None
            return met_data
        except Exception as e:
            st.error(f"Error loading meteorology data: {e}")
            return None

    # Step 2: Plume Rise Calculation
    def calculate_plume_rise(Q, u, stack_diameter, exit_velocity, exit_temperature, ambient_temperature):
        g = 9.81  # Acceleration due to gravity (m/s²)
        buoyancy_flux = (g * stack_diameter**2 * exit_velocity * (exit_temperature - ambient_temperature)) / (
            4 * ambient_temperature
        )
        if u < 2:
            return 21.425 * (buoyancy_flux**(1/3))
        else:
            return 1.6 * (buoyancy_flux / u)**(1/3)

    # Step 3: Gaussian Plume Model Calculation
    def gaussian_plume_corrected(Q, u, H, x, y, sigma_y, sigma_z, deposition_velocity=0):
        term1 = Q / (2 * np.pi * u * sigma_y * sigma_z)
        term2 = np.exp(-y**2 / (2 * sigma_y**2))
        term3 = np.exp(-(x - H)**2 / (2 * sigma_z**2))
        concentration = term1 * term2 * term3

        if deposition_velocity > 0:
            concentration *= np.exp(-deposition_velocity * x / u)
        return concentration

    # Step 4: Dispersion Coefficients
    def pasquill_gifford_sigma_corrected(stability_class, x, mixing_height):
        stability_classes = {
            'A': (0.22, 0.20),
            'B': (0.16, 0.12),
            'C': (0.11, 0.08),
            'D': (0.08, 0.06),
            'E': (0.06, 0.03),
            'F': (0.04, 0.016),
        }
        if stability_class not in stability_classes:
            raise ValueError("Invalid stability class. Use 'A', 'B', 'C', 'D', 'E', or 'F'.")

        ay, az = stability_classes[stability_class]
        sigma_y = ay * x**0.894
        sigma_z = az * x**0.894

        if mixing_height > 0:
            sigma_z = min(sigma_z, 0.5 * mixing_height)
        return sigma_y, sigma_z

    # Step 5: Main Calculation Function
    def calculate_dispersion_2d(met_data, selected_dates, hour_range, Q, stack_height, stack_diameter,
                                exit_velocity, exit_temperature, ambient_temperature, deposition_velocity, center_lat, center_lon):
        """
        Calculate the 2D dispersion model for selected hours of a specific day and month.

        Parameters:
        met_data : DataFrame : Meteorological data
        selected_dates : list : List of tuples (month, day)
        hour_range : tuple : Start and end hour
        Other parameters : Stack and plume properties
        center_lat : float : Center latitude
        center_lon : float : Center longitude

        Returns:
        List : Heatmap data with pollutant concentrations
        """
        filtered_data = met_data[
            (met_data[['MO', 'DY']].apply(tuple, axis=1).isin(selected_dates)) &
            (met_data['HR'] >= hour_range[0]) &
            (met_data['HR'] <= hour_range[1])
        ]
        if filtered_data.empty:
            st.warning("No meteorological data available for the selected filters.")
            return None

        # Grid for 2D dispersion
        x_range = np.linspace(0, 2000, 100)  # Downwind distance (m)
        y_range = np.linspace(-500, 500, 100)  # Crosswind distance (m)
        X, Y = np.meshgrid(x_range, y_range)
        concentrations = np.zeros_like(X)

        # Process each meteorological condition
        for _, row in filtered_data.iterrows():
            u = row['WS10M']
            wind_direction = row['WD10M']
            stability_class = 'D'  # Default to neutral stability
            mixing_height = 1000  # Default mixing height

            plume_rise = calculate_plume_rise(Q, u, stack_diameter, exit_velocity, exit_temperature, ambient_temperature)
            H_eff = stack_height + plume_rise

            for i in range(X.shape[0]):
                for j in range(X.shape[1]):
                    x = X[i, j]
                    y = Y[i, j]

                    sigma_y, sigma_z = pasquill_gifford_sigma_corrected(stability_class, x, mixing_height)
                    concentrations[i, j] += gaussian_plume_corrected(Q, u, H_eff, x, y, sigma_y, sigma_z, deposition_velocity)

        # Convert grid coordinates to geographical points
        heatmap_data = []
        for i in range(X.shape[0]):
            for j in range(X.shape[1]):
                lat = center_lat + (Y[i, j] / 111320)  # Approx conversion of meters to lat
                lon = center_lon + (X[i, j] / (111320 * np.cos(np.radians(center_lat))))
                heatmap_data.append([lat, lon, concentrations[i, j]])

        return heatmap_data

    # Step 6: Plot 2D Dispersion on Folium Map
    def plot_2d_dispersion_on_folium(heatmap_data, center_lat, center_lon, legend):
        # Initialize map centered on the given coordinates
        m = folium.Map(location=[center_lat, center_lon], zoom_start=10)

        # Add heatmap data
        HeatMap(
            [[lat, lon, val] for lat, lon, val in heatmap_data if val > 0],  # Filter non-zero concentrations
            radius=10,
            blur=15,
            max_zoom=10
        ).add_to(m)

        # Add legend on the right
        st.sidebar.markdown("### Legend: Pollutant Concentration (g/m³)")
        for label, color in legend.items():
            st.sidebar.markdown(f"<span style='color:{color}'>{label}</span>", unsafe_allow_html=True)

        return m

    # Upload Meteorology Data
    csv_file = st.file_uploader("Upload your meteorology CSV file", type=["csv"])
    if csv_file:
        met_data = pd.read_csv(csv_file)
        st.success("Meteorology data loaded successfully!")

        # User inputs
        selected_dates = st.multiselect("Select Dates (Month, Day):", met_data[['MO', 'DY']].drop_duplicates().apply(tuple, axis=1).tolist())
        hour_start = st.number_input("Enter the start hour (0-23):", min_value=0, max_value=23)
        hour_end = st.number_input("Enter the end hour (0-23):", min_value=0, max_value=23)
        hour_range = (hour_start, hour_end)

        Q = st.number_input("Emission rate (g/s):")
        stack_height = st.number_input("Stack height (m):")
        stack_diameter = st.number_input("Stack diameter (m):")
        exit_velocity = st.number_input("Exhaust gas exit velocity (m/s):")
        exit_temperature = st.number_input("Exhaust gas temperature (K):")
        ambient_temperature = st.number_input("Ambient air temperature (K):")
        deposition_velocity = st.number_input("Deposition velocity (m/s):")

        center_lat = st.number_input("Enter the latitude of the area center:")
        center_lon = st.number_input("Enter the longitude of the area center:")

        # Legend for EPA
        legend = {
            "Low Concentration": "green",
            "Moderate Concentration": "yellow",
            "High Concentration": "red"
        }

        # Calculate and plot
        if st.button("Run Plume Dispersion Model"):
            heatmap_data = calculate_dispersion_2d(met_data, selected_dates, hour_range, Q, stack_height,
                                                   stack_diameter, exit_velocity, exit_temperature, ambient_temperature,
                                                   deposition_velocity, center_lat, center_lon)
            if heatmap_data:
                map_dispersion = plot_2d_dispersion_on_folium(heatmap_data, center_lat, center_lon, legend)
                folium_static(map_dispersion)





# About section
if options == "About":
    st.subheader("📄 About CEC (Clanert Enviro Cloud)")
    st.write("""
CEC is an interactive model designed for:
- **Meteorological wind data analysis**, including wind rose plots, wind class distribution, and geographical overlays.
- **Air quality analysis**, with features for calculating AQI, visualizing pollutant trends, and generating comprehensive reports.


---

### About Clanert Sustain
Clanert Sustain is a specialized environmental company and solutions provider, committed to addressing the growing challenges of pollution, waste management, and climate change. Through innovative strategies, cutting-edge technology, and engineering expertise, we deliver impactful solutions that transform industries and communities.

We collaborate with businesses, government agencies, and communities to create sustainable systems that contribute to a cleaner, greener planet.

Our work is guided by principles of **sustainability**, **affordability**, and **collaboration**. By focusing on:
- **Environmental monitoring**
- **Waste management**
- **Renewable energy**
- **Pollution mitigation**

We align our efforts with both global standards and local needs, ensuring a positive environmental and societal impact.

---

### Contact Us
📍 **Address**: Mbezi, Dar es Salaam, Tanzania  
📧 **Email**: [info@clanertsustain.co.tz](mailto:info@clanertsustain.co.tz)  
📞 **Phone**: +255 621 667 315  
🌐 **Website**: [www.clanertsustain.co.tz](http://www.clanertsustain.co.tz)

© 2024 **Clanert Sustain**. All rights reserved.
""")
